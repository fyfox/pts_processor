from docx import Document
import xlwings as xw

class PTS_Processor(object):
    def __init__(self, doc_path):
        f = open(doc_path, 'rb')
        self.doc= Document(f)
        self.cmd_xls_name = doc_path.split('.')[0]+'.xlsx'
        self.cmd_list=[]

    def read_cmd(self):
        for n in range(len(self.doc.tables)):
            row_title = [r.text for r in self.doc.tables[n].row_cells(0)]

            if row_title[0].find("ID") is not -1:
                row_cmd = [r.text for r in self.doc.tables[n].row_cells(1)]
                if row_cmd[0] == "":
                        row_cmd = [r.text for r in self.doc.tables[n].row_cells(2)]
                #print(row_title)
                #print(row_cmd)
                cmd_str = " ".join(row_cmd[1:])
                self.cmd_list.append([row_cmd[0],cmd_str])

    def create_cmd_xlsx(self):
        print(self.cmd_xls_name)
        #method 1
        #app=xw.App(visible=True,add_book=False)
        #self.wb=app.books.add()

        #method 2
        self.wb = xw.Book()
        self.sht = self.wb.sheets[0]


    def write_cmd_xlsx(self):
        #self.wb.sheets[0].range('A1:D5')
        rng_str = 'B1:B'+str(len(self.cmd_list))
        rng = self.sht.range(rng_str)
        rng.value = self.cmd_list

    def close_cmd_xlsx(self):
        self.wb.save(self.cmd_xls_name)
        self.wb.close()
        print("close xls")

if __name__ == "__main__":
    file = "ESTONA_EL-ET-CAR_ID5027-88_V01_20161130-Release.docx"
    pts_or = PTS_Processor(file)
    pts_or.read_cmd()
    pts_or.create_cmd_xlsx()
    pts_or.write_cmd_xlsx()
    pts_or.close_cmd_xlsx()

