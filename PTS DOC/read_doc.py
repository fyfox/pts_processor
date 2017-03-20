
from docx import Document

def read_cmd(doc):
    for n in range(len(doc.tables)):
        row_title = [r.text for r in doc.tables[n].row_cells(0)]

        if row_title[0].find("ID") is not -1:
            row_cmd = [r.text for r in doc.tables[n].row_cells(1)]
            if row_cmd[0] == "":
                    row_cmd = [r.text for r in doc.tables[n].row_cells(2)]
            print(row_title)
            print(row_cmd)

if __name__ == "__main__":
    print("ok")
    f = open("ESTONA_EL-ET-CAR_ID5027-88_V01_20161130-Release.docx", 'rb')
    pts = Document(f)
    read_cmd(pts)
